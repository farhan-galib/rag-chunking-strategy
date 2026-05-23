"""
File extraction utilities for multiple document formats
Supports: TXT, MD, PDF, DOCX, XLSX
"""

import io
from typing import Any, Dict, Optional, Tuple
from pathlib import Path

def extract_text_from_file(file_content: bytes, file_extension: str) -> Tuple[str, bool]:
    """
    Extract text from any supported file format.
    """
    file_extension = file_extension.lower().strip('.')
    
    try:
        if file_extension in ['txt', 'md']:
            try:
                return file_content.decode('utf-8'), True
            except:
                return file_content.decode('latin-1'), True
        
        elif file_extension == 'pdf':
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text.strip(), True
            
        elif file_extension in ['docx', 'doc']:
            import docx
            doc = docx.Document(io.BytesIO(file_content))
            text = "\n".join([para.text for para in doc.paragraphs])
            return text.strip(), True
            
        elif file_extension in ['xlsx', 'xls']:
            import pandas as pd
            df = pd.read_excel(io.BytesIO(file_content))
            return df.to_string(), True
            
        return "Unsupported format", False
    except Exception as e:
        return f"Extraction error: {str(e)}", False

def get_document_metadata(file_content: bytes, file_extension: str, filename: str) -> Dict[str, Any]:
    """
    Extract display metadata: title, page count, and formatted size.
    """
    size_bytes = len(file_content)
    title = Path(filename).stem.replace("_", " ").replace("-", " ").strip() or filename
    file_extension = file_extension.lower().strip('.')
    
    page_count = 1
    
    try:
        if file_extension == 'pdf':
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_content))
            page_count = len(reader.pages)
        elif file_extension in ['xlsx', 'xls']:
            import pandas as pd
            engine = 'openpyxl' if file_extension == 'xlsx' else None
            xl = pd.ExcelFile(io.BytesIO(file_content), engine=engine)
            page_count = len(xl.sheet_names)
        elif file_extension in ['docx', 'doc']:
            try:
                import docx
                doc = docx.Document(io.BytesIO(file_content))
                # python-docx core_properties.pages is often 0 or 1 unless fields are updated.
                # Word-based estimation is generally more reliable for business docs.
                # ~450 words per page is a standard average.
                full_text = "\n".join([p.text for p in doc.paragraphs])
                words = len(full_text.split())
                page_count = max(1, round(words / 450))
                
                # Check if we can get a better count from metadata if it exists
                cp = doc.core_properties
                if cp.pages and cp.pages > 1:
                    page_count = cp.pages
            except:
                page_count = 1
        elif file_extension in ['txt', 'md']:
            try:
                text = file_content.decode('utf-8', errors='ignore')
                # Estimate ~500 words per page for plain text
                words = len(text.split())
                page_count = max(1, round(words / 500))
            except:
                page_count = 1
    except:
        page_count = 1

    return {
        "title": title,
        "page_count": page_count, 
        "size_kb": f"{size_bytes / 1024:.1f} KB",
        "size_bytes": size_bytes,
    }

def get_supported_file_types() -> list:
    """
    Get list of supported file extensions.
    """
    return ['txt', 'md', 'pdf', 'docx', 'xlsx', 'xls']

def docx_to_html(content: bytes) -> str:
    """Convert DOCX to simple HTML."""
    import docx
    doc = docx.Document(io.BytesIO(content))
    html = ""
    for para in doc.paragraphs:
        if para.text.strip():
            html += f"<p>{para.text}</p>"
    return html

def load_excel_sheets(content: bytes, ext: str) -> Dict:
    """Load all sheets from an Excel file."""
    import pandas as pd
    engine = 'openpyxl' if ext == 'xlsx' else None
    xl = pd.ExcelFile(io.BytesIO(content), engine=engine)
    sheets = {}
    for name in xl.sheet_names:
        sheets[name] = xl.parse(name)
    return {
        "sheet_names": xl.sheet_names,
        "sheets": sheets
    }
